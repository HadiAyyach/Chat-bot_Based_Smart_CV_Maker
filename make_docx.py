from docx import Document
from docx.shared import Inches,RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import re

# opening the optimized data
with open('optimized_data.json', 'r') as file:
    optimized_data = json.load(file)

def turn_to_clean_array(data):
    output = []
    for one_data in data :
        data_array = [line for line in  re.split(r"[\n\t\*\+]+", one_data) if line.strip()]
        output.append(data_array) 
    return output

def name(text):
    heading = document.add_heading()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = heading.add_run(text, 0)
    name.font.color.rgb = RGBColor(0, 0, 0)
    name.font.size = Pt(28)
    name.bold = True

def personal_info(text):
    personal_info = document.add_paragraph()
    personal_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_i = personal_info.add_run(text)
    p_i.bold = True
    p_i.font.size = Pt(11)

def title(text):
    title = document.add_heading(level=1)
    tit = title.add_run(text)
    tit.font.color.rgb = RGBColor(0, 0, 0)
    tit.font.size = Pt(15)
    # line = title.add_run("\n______________________________________________________________________________________________________________")
    # line.font.color.rgb = RGBColor(0, 0, 0)
    # line.bold = True
    # line.font.size =Pt(8)
    paragraph2 = document.add_paragraph("________________________________________________________________________________________________________")
    paragraph2.paragraph_format.line_spacing = Pt(1)  


    
    
def details(text,bald = False ):
    if "\n" in text[:3]:
        text = text[1:]
    details = document.add_paragraph()
    det = details.add_run(text)
    det.font.size = Pt(11)
    if bald:
        det.font.bold = True
    # table = document.add_table(rows=1, cols=1)
    # table.cell(0,0).text = text
    # table.cell(0,0).paragraphs[0].runs[0].font.size = Pt(11)
    # if bald:
    #     table.cell(0,0).paragraphs[0].runs[0].font.bold = True


# Create a new document
document = Document()

# Add name and contact information
name(optimized_data["full_name"])
personal_info(optimized_data["email"] + " | " + optimized_data["contact_info"])

# add the summary 
title("SUMMARY")
details(optimized_data["summary"])

# Add work experience
title("EXPERIENCE")
experience = turn_to_clean_array(optimized_data["experience"])
for exp in experience:
    details(exp[0],True)
    all = ""
    for line in exp[1:]:
        all += "\n*" + line 
    details( all )
   

# Add education
title("EDUCATION")
education = turn_to_clean_array(optimized_data["education"]) 
all = ""
for edu in education:
    all += "*" + edu[0] + "\n " 
details(all,True)

# Add skills
title("SKILLS")
skills = turn_to_clean_array(optimized_data["skills"])

all = ""
for skill in skills[0]:
    all += "*" + skill + "\n " 
details(all)

# Add additional information
title("ADDITIONAL INFORMATION  ")
add_info = turn_to_clean_array(optimized_data["additional_info"])

all = ""
for info in add_info[0]:
    all +=  info + "\n " 
details(all)

# Save the document
document.save("ats-friendly-cv.docx")

from docx2pdf import convert

# turn to pdf
docx_file_path = "ats-friendly-cv.docx"
pdf_file_path = "ats-friendly-cv.pdf"
convert(docx_file_path, pdf_file_path)
